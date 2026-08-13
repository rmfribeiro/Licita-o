from __future__ import annotations
import io
import pdfplumber
from docx import Document

# -----------------------------------------------------------------------------
# LIMITE DE TEXTO NA LEITURA DOS ARQUIVOS
# -----------------------------------------------------------------------------
# Este é o corte MAIS A MONTANTE do sistema: onze pontos do app leem documentos
# por aqui, e o que for cortado neste ponto não chega a módulo nenhum.
#
# Estava em 50.000 caracteres. Medido em 13/08/2026 com um Termo de Referência
# real de obra (40 páginas, 136.556 caracteres): 63% do documento era descartado
# ANTES da análise. Pior, a correção feita nos módulos de IA no dia anterior
# (limite de 300.000 em ia_utils.preparar_documento) não adiantava nada para
# quem passa por aqui — o texto já chegava mutilado.
#
# 300.000 caracteres acompanham o limite de ia_utils.LIMITE_DOC_PADRAO: é o que
# cabe com folga na janela do modelo. Praticamente todo edital, TR ou contrato
# brasileiro entra inteiro.
_LIMITE_CHARS = 300_000


def _extrair_pdf(conteudo: bytes) -> str:
    texto = ""
    try:
        with pdfplumber.open(io.BytesIO(conteudo)) as pdf:
            for page in pdf.pages:
                texto += page.extract_text() or ""
    except Exception:
        pass
    return texto


def _extrair_docx(conteudo: bytes) -> str:
    texto = ""
    try:
        doc = Document(io.BytesIO(conteudo))
        for para in doc.paragraphs:
            if para.text.strip():
                texto += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        texto += cell.text + "\n"
    except Exception:
        pass
    return texto


def extrair_texto(arquivos: list) -> tuple[str, list[str]]:
    partes: list[str] = []
    avisos: list[str] = []

    for arquivo in arquivos:
        nome = arquivo.name
        if hasattr(arquivo, "seek"):
            arquivo.seek(0)
        conteudo = arquivo.read() if hasattr(arquivo, "read") else arquivo.getvalue()
        ext = nome.lower().rsplit(".", 1)[-1] if "." in nome else ""

        if ext == "pdf":
            texto = _extrair_pdf(conteudo)
        elif ext == "docx":
            texto = _extrair_docx(conteudo)
        else:
            avisos.append(f"Formato não suportado ignorado: {nome}")
            continue

        if not texto.strip():
            avisos.append(f"Sem texto extraível: {nome}")
            continue

        partes.append(f"[ARQUIVO: {nome}]\n{texto.strip()}")

    if not partes:
        raise ValueError("Nenhum texto extraível nos arquivos enviados.")

    concatenado = "\n\n".join(partes)

    if len(concatenado) > _LIMITE_CHARS:
        _original = len(concatenado)
        concatenado = concatenado[:_LIMITE_CHARS]
        # O aviso antigo dizia apenas "documentos extensos podem ter conteúdo
        # não analisado" — vago demais para quem precisa decidir se confia no
        # parecer. Agora informa o tamanho real, quanto ficou de fora e o que
        # fazer, porque o que não foi lido não foi auditado.
        _fora = _original - _LIMITE_CHARS
        # Formata o separador de milhar SÓ nos números. Aplicar .replace(",", ".")
        # na frase inteira comia as vírgulas do texto ("analise cada uma. ou envie").
        _br = lambda n: f"{n:,}".replace(",", ".")
        avisos.append(
            f"ATENÇÃO: os documentos somam {_br(_original)} caracteres e a leitura foi "
            f"limitada a {_br(_LIMITE_CHARS)} — {_br(_fora)} caracteres "
            f"({_fora/_original:.0%} do total) NÃO foram analisados. O que não foi lido "
            "não foi auditado: divida o material em partes menores e analise cada uma, "
            "ou envie separadamente os documentos mais relevantes."
        )

    return concatenado, avisos
