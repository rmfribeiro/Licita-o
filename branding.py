# -*- coding: utf-8 -*-
"""Timbre configuravel para os pareceres. Edite branding.json com os dados da
sua empresa (nome, contato, cor e, opcionalmente, o caminho de um logo PNG)."""
import json, os, re
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_AQUI = os.path.dirname(os.path.abspath(__file__))
_PADRAO = {"empresa": "[SUA EMPRESA]", "tagline": "[SLOGAN DA EMPRESA]",
           "referencia_legal": "Auditoria de Editais — Lei 14.133/2021",
           "contato": "", "cor_primaria": "1F4E79", "logo": "",
           # logo_timbre: versao com fundo branco para documentos (PDF/DOCX
           # lidam mal com PNG transparente); icone: monograma quadrado do
           # favicon. Ambos opcionais — sem eles, cai no "logo".
           "logo_timbre": "", "icone": ""}


def caminho(nome_chave: str) -> str:
    """Devolve o caminho absoluto da imagem da chave pedida, ou "" se nao
    existir no disco. Aceita cair para "logo" quando a chave especifica
    (logo_timbre/icone) nao foi configurada."""
    b = carregar()
    for chave in (nome_chave, "logo"):
        arq = b.get(chave)
        if arq:
            p = os.path.join(_AQUI, arq)
            if os.path.isfile(p):
                return p
    return ""

def carregar():
    p = os.path.join(_AQUI, "branding.json")
    if os.path.isfile(p):
        try:
            with open(p, encoding="utf-8") as fh:
                data = json.load(fh)
            # Normaliza None para o padrão; valida cor_primaria como hex de 6 dígitos
            data = {k: (v if v is not None else _PADRAO.get(k, "")) for k, v in data.items()}
            cor = str(data.get("cor_primaria", "")).lstrip("#")
            data["cor_primaria"] = cor if re.fullmatch(r"[0-9A-Fa-f]{6}", cor) else _PADRAO["cor_primaria"]
            return {**_PADRAO, **data}
        except Exception:
            pass
    return dict(_PADRAO)

def cabecalho_pdf(estilo_titulo, altura_mm: float = 26.0):
    """Cabecalho de marca para os relatorios em PDF (ReportLab).

    Devolve a lista de flowables a ser inserida no topo do documento:
    a logomarca centralizada (quando o arquivo existe) seguida do nome da
    empresa em corpo menor; sem o arquivo, cai no titulo em texto de sempre,
    para que nenhum relatorio deixe de ser gerado por falta de imagem.

    Importa o ReportLab aqui dentro (e nao no topo do modulo) porque
    branding.py tambem e usado pelos geradores .docx, que nao dependem dele.
    """
    from reportlab.platypus import Image, Spacer
    from reportlab.lib.units import mm
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import Paragraph
    from reportlab.lib import colors

    b = carregar()
    _rotulo = f"{b.get('produto', 'RM Lisura')} — {b['empresa']}"
    _arq = caminho("logo_timbre")
    if not _arq:
        return [Paragraph(_rotulo, estilo_titulo)]

    try:
        from PIL import Image as _PILImage
        with _PILImage.open(_arq) as _im:
            _larg_px, _alt_px = _im.size
        altura = altura_mm * mm
        largura = altura * (_larg_px / float(_alt_px))
        img = Image(_arq, width=largura, height=altura)
        img.hAlign = "CENTER"
        _estilo_empresa = ParagraphStyle(
            "marca_empresa", parent=estilo_titulo,
            fontSize=9.5, leading=11, spaceBefore=2, spaceAfter=2,
            textColor=colors.HexColor("#" + b["cor_primaria"]),
        )
        # A logomarca ja diz "RM LISURA"; repetir o nome do produto ao lado
        # polui o topo. Fica so a razao social, que a imagem nao traz.
        return [img, Spacer(1, 2), Paragraph(b["empresa"], _estilo_empresa)]
    except Exception:
        return [Paragraph(_rotulo, estilo_titulo)]


def add_banner(doc):
    """Insere o timbre (logo opcional + nome da empresa + tagline) no topo."""
    b = carregar()
    cor = RGBColor.from_string(b["cor_primaria"])
    _logo = caminho("logo_timbre")
    if _logo:
        try:
            pic = doc.add_paragraph(); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic.add_run().add_picture(_logo, height=Mm(16))
            pic.paragraph_format.space_after = Pt(2)
        except Exception:
            pass
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(b["empresa"]); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = cor
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(8)
    r2 = p2.add_run(b["referencia_legal"]); r2.font.size = Pt(8.5); r2.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    pPr = p2._p.get_or_add_pPr(); pb = OxmlElement("w:pBdr"); bt = OxmlElement("w:bottom")
    bt.set(qn("w:val"), "single"); bt.set(qn("w:sz"), "4"); bt.set(qn("w:space"), "4"); bt.set(qn("w:color"), b["cor_primaria"])
    pb.append(bt); pPr.append(pb)
    return b

def add_contato_footer(sec):
    """Escreve a linha de contato no rodape, preservando paragrafos com numeracao de pagina."""
    b = carregar()
    if not b.get("contato"):
        return
    footer = sec.footer
    # Reutiliza o ultimo paragrafo sem campo de pagina; senao adiciona novo
    p = None
    for para in reversed(footer.paragraphs):
        if not para._p.findall(".//" + qn("w:fldChar")):
            p = para
            break
    if p is None:
        p = footer.add_paragraph()
    for r in list(p._p.findall(".//" + qn("w:r"))):
        r.getparent().remove(r)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(b["empresa"] + "   ·   " + b["contato"])
    r.font.size = Pt(7.5); r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
